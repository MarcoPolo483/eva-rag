"""DOCX document loader using python-docx."""
from typing import BinaryIO

from docx import Document

from eva_rag.loaders.base import DocumentLoader, ExtractedDocument


class DOCXLoader(DocumentLoader):
    """Load and extract text from DOCX files."""
    
    def load(self, file: BinaryIO, filename: str) -> ExtractedDocument:
        """
        Extract text from DOCX file.
        
        Args:
            file: Binary DOCX file object
            filename: Original filename
            
        Returns:
            Extracted document with text and paragraph count
            
        Raises:
            ValueError: If DOCX is invalid or empty
        """
        try:
            # Read DOCX
            document = Document(file)
            
            if not document.paragraphs:
                raise ValueError(f"DOCX file '{filename}' contains no paragraphs")
            
            # Extract text from all paragraphs
            text_parts: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            if not text_parts:
                raise ValueError(f"DOCX file '{filename}' contains no extractable text")
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata from core properties
            metadata = {}
            if document.core_properties.title:
                metadata["title"] = document.core_properties.title
            if document.core_properties.author:
                metadata["author"] = document.core_properties.author
            
            return ExtractedDocument(
                text=full_text,
                page_count=len(document.paragraphs),  # Use paragraph count as proxy
                metadata=metadata,
            )
            
        except Exception as e:
            raise ValueError(f"Failed to load DOCX '{filename}': {str(e)}") from e
